#!/usr/bin/env python3
"""
词表提取与分析脚本
从 ~/stu/wordbank/word/ 提取全部英文单词，按教育阶段分类，计算重叠。
"""

import os, re, json
from collections import OrderedDict, defaultdict

# ─── 路径配置 ───
BASE = os.path.expanduser("~/stu/wordbank/word/")
OUT_DIR = os.path.expanduser("~/stu/vocab_estimator/data/")
os.makedirs(OUT_DIR, exist_ok=True)

# ─── 各阶段优先级定义 ───
STAGE_PRIORITY = OrderedDict([
    ("primary_3",   {"priority": 1,  "label": "小学三年级"}),
    ("primary_4",   {"priority": 2,  "label": "小学四年级"}),
    ("primary_5",   {"priority": 3,  "label": "小学五年级"}),
    ("primary_6",   {"priority": 4,  "label": "小学六年级"}),
    ("junior_7",    {"priority": 5,  "label": "七年级"}),
    ("junior_8",    {"priority": 6,  "label": "八年级"}),
    ("junior_9",    {"priority": 7,  "label": "九年级"}),
    ("senior",      {"priority": 8,  "label": "高中"}),
    ("cet4",        {"priority": 9,  "label": "大学四级"}),
    ("cet6",        {"priority": 10, "label": "大学六级"}),
    ("ielts",       {"priority": 11, "label": "雅思"}),
])


def clean_word(raw):
    """清洗单个英文单词：去空格/标点/中文/数字，统一小写"""
    if not raw:
        return None
    text = str(raw).strip()
    # 去 \xa0 等不可见字符
    text = text.replace('\xa0', ' ')
    text = text.strip()
    if not text:
        return None
    # 过滤纯中文/纯数字/空行
    text_lower = text.lower()

    # 短语情况：允许空格和连字符
    # 检查是否含有英文字母
    if not re.search(r'[a-z]', text_lower):
        return None

    # 分词处理
    # 先按空格拆分，对每个部分清洗
    parts = text_lower.split()
    cleaned_parts = []
    for part in parts:
        # 去掉两端非字母数字字符（保留连字符、撇号）
        part = re.sub(r'^[^a-z0-9\'-]+|[^a-z0-9\'-]+$', '', part)
        if part:
            cleaned_parts.append(part)

    if not cleaned_parts:
        return None

    return ' '.join(cleaned_parts)


def extract_from_primary_xlsx(path):
    """
    人教版PEP小学英语三年级起点单词表
    Col A: 课本(三上/三下/四上...), Col B: Unit, Col C: 单词
    返回 {stage_name: [word, ...]}
    """
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True)
    ws = wb.active
    stage_map = {
        '三上': 'primary_3', '三下': 'primary_3',
        '四上': 'primary_4', '四下': 'primary_4',
        '五上': 'primary_5', '五下': 'primary_5',
        '六上': 'primary_6', '六下': 'primary_6',
    }
    stage_words = defaultdict(set)
    issues = []

    for row in ws.iter_rows(min_row=2, values_only=True):
        if len(row) < 3:
            continue
        book = str(row[0]).strip() if row[0] else ''
        word_raw = str(row[2]).strip() if row[2] else ''
        stage = stage_map.get(book)
        if not stage:
            if book not in ('课本', 'None', ''):
                issues.append(f"未知课本列值: {book}, 单词: {word_raw}")
            continue
        cleaned = clean_word(word_raw)
        if cleaned:
            stage_words[stage].add(cleaned)
        elif word_raw and word_raw.strip() and word_raw not in ('课本', '单词', '', None):
            issues.append(f"未能清洗: col=({book}) word={repr(word_raw)}")

    wb.close()
    return stage_words, issues


def extract_from_junior_table(path, stage_label):
    """
    初中 docx 文件：表格中 col index 2 = 英文
    返回 {stage_name: [word, ...]}
    """
    from docx import Document
    doc = Document(path)
    stage_words = defaultdict(set)
    issues = []

    if doc.tables:
        table = doc.tables[0]
        for row in table.rows:
            cells = row.cells
            if len(cells) < 3:
                continue
            word_raw = cells[2].text.strip()
            cleaned = clean_word(word_raw)
            if cleaned:
                stage_words[stage_label].add(cleaned)
            elif word_raw and word_raw not in ('单词', '', '词汇', '英语') and len(word_raw) > 1:
                issues.append(f"未能清洗: word={repr(word_raw)}")
    else:
        issues.append(f"无表格: {path}")

    return stage_words, issues


def extract_from_gaokao_xls(path):
    """
    高考乱序版3500词汇.xls
    Col 0=序号, Col 1=单词, Col 2=音标, Col 3=中文
    返回 {'senior': [word, ...]}
    """
    import xlrd
    wb = xlrd.open_workbook(path)
    ws = wb.sheet_by_index(0)
    stage_words = defaultdict(set)
    issues = []

    for r in range(1, ws.nrows):
        word_raw = str(ws.cell_value(r, 1)).strip() if ws.ncols > 1 else ''
        if not word_raw or word_raw == '单词':
            continue
        cleaned = clean_word(word_raw)
        if cleaned:
            stage_words['senior'].add(cleaned)
        elif word_raw not in ('', '单词', '序号'):
            issues.append(f"未能清洗:Row{r} word={repr(word_raw)}")

    return stage_words, issues


def extract_from_cet_xls(path, stage_label):
    """
    四六级词汇 xls 文件
    Col 0=序号, Col 1=单词, Col 2=音标, Col 3=释义
    返回 {stage_label: [word, ...]}
    """
    import xlrd
    wb = xlrd.open_workbook(path)
    ws = wb.sheet_by_index(0)
    stage_words = defaultdict(set)
    issues = []

    for r in range(1, ws.nrows):
        word_raw = str(ws.cell_value(r, 1)).strip() if ws.ncols > 1 else ''
        if not word_raw or word_raw == '单词':
            continue
        cleaned = clean_word(word_raw)
        if cleaned:
            stage_words[stage_label].add(cleaned)
        elif word_raw not in ('', '单词', '序号'):
            issues.append(f"未能清洗:Row{r} word={repr(word_raw)}")

    return stage_words, issues


def extract_from_ielts_xls(path, stage_label):
    """
    雅思词汇 xls 文件
    Col 0=序号, Col 1=单词, Col 2=释义
    返回 {stage_label: [word, ...]}
    """
    import xlrd
    wb = xlrd.open_workbook(path)
    ws = wb.sheet_by_index(0)
    stage_words = defaultdict(set)
    issues = []

    for r in range(1, ws.nrows):
        word_raw = str(ws.cell_value(r, 1)).strip() if ws.ncols > 1 else ''
        if not word_raw:
            continue
        # 跳过类似 'A.D.', 'A.M.' 这种纯大写缩写
        if re.match(r'^[A-Z.]+$', word_raw.strip()):
            continue
        cleaned = clean_word(word_raw)
        if cleaned:
            stage_words[stage_label].add(cleaned)
        elif word_raw not in ('', '单词', '序号'):
            issues.append(f"未能清洗:Row{r} word={repr(word_raw)}")

    return stage_words, issues


def process_file(path, file_key, parser_func, *parser_args):
    """安全处理单个文件的封装"""
    print(f"\n{'='*60}")
    print(f"处理文件: {path}")
    print(f"标签: {file_key}")
    try:
        stage_words, issues = parser_func(path, *parser_args)
        total = sum(len(v) for v in stage_words.values())
        print(f"结果: {total} 个去重单词")
        for stage, words in sorted(stage_words.items()):
            sample = sorted(words)[:20]
            print(f"  {stage}: {len(words)} 词, 样例: {sample}")
        if issues:
            print(f"  ⚠ 问题数: {len(issues)}")
            for iss in issues[:10]:
                print(f"    - {iss}")
        return stage_words, issues, total
    except Exception as e:
        print(f"  ❌ 错误: {e}")
        import traceback
        traceback.print_exc()
        return defaultdict(set), [str(e)], 0


# ══════════════════════════════════════════════════════════════
# 步骤 1 和 2：提取 + 阶段映射
# ══════════════════════════════════════════════════════════════

print("╔════════════════════════════════════════════════════════╗")
print("║         词表提取与分析 — Step 1 & 2                    ║")
print("╚════════════════════════════════════════════════════════╝")

all_issues = {}
# 主词典：stage_name -> set(words)
all_stage_words = defaultdict(set)

# ---- 1. 小学 ----
path_primary = os.path.join(BASE, "人教版PEP小学英语三年级起点单词表(三年级至六年级全8册) 共20页.xlsx")
sw, iss, cnt = process_file(path_primary, "小学", extract_from_primary_xlsx)
all_issues["primary"] = iss
for stage, words in sw.items():
    all_stage_words[stage].update(words)

# ---- 2. 七年级 ----
junior_base = os.path.join(BASE, "人教版初中英语单词短语默写表/分册/")
path_j7a = os.path.join(junior_base, "人教版七年级上册英语单词英译汉.docx")
sw, iss, cnt = process_file(path_j7a, "七年级上册", extract_from_junior_table, "junior_7")
all_issues["七年级上册"] = iss
for stage, words in sw.items():
    all_stage_words[stage].update(words)

path_j7b = os.path.join(junior_base, "人教版七年级下册英语单词英译汉.docx")
sw, iss, cnt = process_file(path_j7b, "七年级下册", extract_from_junior_table, "junior_7")
all_issues["七年级下册"] = iss
for stage, words in sw.items():
    all_stage_words[stage].update(words)

# ---- 3. 八年级 ----
path_j8a = os.path.join(junior_base, "人教版八年级上册英语单词英译汉.docx")
sw, iss, cnt = process_file(path_j8a, "八年级上册", extract_from_junior_table, "junior_8")
all_issues["八年级上册"] = iss
for stage, words in sw.items():
    all_stage_words[stage].update(words)

path_j8b = os.path.join(junior_base, "人教版八年级下册英语单词英译汉.docx")
sw, iss, cnt = process_file(path_j8b, "八年级下册", extract_from_junior_table, "junior_8")
all_issues["八年级下册"] = iss
for stage, words in sw.items():
    all_stage_words[stage].update(words)

# ---- 4. 九年级 ----
path_j9 = os.path.join(junior_base, "最新人教版九年级英语全一册单词默写板（英译汉）.doc")
sw, iss, cnt = process_file(path_j9, "九年级全一册", extract_from_junior_table, "junior_9")
all_issues["九年级全一册"] = iss
for stage, words in sw.items():
    all_stage_words[stage].update(words)

# ---- 5. 高中 ----
path_senior = os.path.join(BASE, "高考乱序版3500词汇.xls")
sw, iss, cnt = process_file(path_senior, "高中", extract_from_gaokao_xls)
all_issues["高中"] = iss
for stage, words in sw.items():
    all_stage_words[stage].update(words)

# ---- 6. 四级 ----
path_cet4 = os.path.join(BASE, "四六级词汇/英语四级词汇/01.大学英语四级词汇完整带音标-可打印-可编辑-正序版.xls")
sw, iss, cnt = process_file(path_cet4, "四级", extract_from_cet_xls, "cet4")
all_issues["四级"] = iss
for stage, words in sw.items():
    all_stage_words[stage].update(words)

# ---- 7. 六级 ----
path_cet6 = os.path.join(BASE, "四六级词汇/英语六级词汇/大学英语六级词汇完整带音标-可打印-可编辑-正序版.xls")
sw, iss, cnt = process_file(path_cet6, "六级", extract_from_cet_xls, "cet6")
all_issues["六级"] = iss
for stage, words in sw.items():
    all_stage_words[stage].update(words)

# ---- 8. 雅思 ----
path_ielts8k = os.path.join(BASE, "06.雅思词汇/雅思8000词/雅思词汇8000词EXCEL版-顺序版.xls")
sw, iss, cnt = process_file(path_ielts8k, "雅思8000", extract_from_ielts_xls, "ielts")
all_issues["雅思8000"] = iss
for stage, words in sw.items():
    all_stage_words[stage].update(words)

path_ielts9k = os.path.join(BASE, "06.雅思词汇/雅思9400词/雅思词汇9400词EXCEL版-顺序版.xls")
sw, iss, cnt = process_file(path_ielts9k, "雅思9400", extract_from_ielts_xls, "ielts")
all_issues["雅思9400"] = iss
for stage, words in sw.items():
    all_stage_words[stage].update(words)


# ══════════════════════════════════════════════════════════════
# 步骤 3：重叠处理与阶段分配
# ══════════════════════════════════════════════════════════════

print("\n\n╔════════════════════════════════════════════════════════╗")
print("║         重叠分析 — Step 3                             ║")
print("╚════════════════════════════════════════════════════════╝")

# 固定阶段的处理顺序（优先级=priority）
stages_order = list(STAGE_PRIORITY.keys())

# word -> 该词出现的 stage 集合
word_appearances = defaultdict(set)
for stage in stages_order:
    for word in all_stage_words[stage]:
        word_appearances[word].add(stage)

# 每个词分配首次出现阶段
word_first_stage = {}
word_first_priority = {}
word_all_stages = defaultdict(list)

for word, stages in word_appearances.items():
    # 按优先级排序
    sorted_stages = sorted(stages, key=lambda s: STAGE_PRIORITY[s]["priority"])
    first_stage = sorted_stages[0]
    first_priority = STAGE_PRIORITY[first_stage]["priority"]
    word_first_stage[word] = first_stage
    word_first_priority[word] = first_priority
    word_all_stages[word] = sorted_stages

# 各阶段独有词数（仅在该阶段首次出现）
stage_exclusive = {}
for stage in stages_order:
    exclusive = {w for w in word_first_stage if word_first_stage[w] == stage}
    stage_exclusive[stage] = exclusive
    print(f"  {stage} ({STAGE_PRIORITY[stage]['label']}): 独有 {len(exclusive)} 词")

# 各阶段累计词数
print("\n--- 累计词数（含更低阶段）---")
cumulative = set()
stage_cumulative = {}
for stage in stages_order:
    cumulative.update(all_stage_words.get(stage, set()))
    stage_cumulative[stage] = len(cumulative)
    print(f"  {stage}: 累计 {len(cumulative)} 词")

# 重叠矩阵
print("\n--- 重叠矩阵 (行=低阶段, 列=高阶段) ---")
n = len(stages_order)
overlap_matrix = {}
for i, low_stage in enumerate(stages_order):
    low_words = all_stage_words.get(low_stage, set())
    row = {}
    for j, high_stage in enumerate(stages_order):
        if j <= i:
            row[high_stage] = None  # 不显示自己或更低
        else:
            high_words = all_stage_words.get(high_stage, set())
            overlap = len(low_words & high_words)
            row[high_stage] = overlap
    overlap_matrix[low_stage] = row
    print(f"  {low_stage:12s}: ", end="")
    for high_stage in stages_order:
        val = row.get(high_stage)
        if val is not None:
            print(f"{high_stage:10s}={val:5d}", end="  ")
    print()


# ══════════════════════════════════════════════════════════════
# 步骤 5：输出 JSON
# ══════════════════════════════════════════════════════════════

print("\n\n╔════════════════════════════════════════════════════════╗")
print("║         输出 JSON  — Step 5                           ║")
print("╚════════════════════════════════════════════════════════╝")

# 构建 JSON 结构
stages_json = {}
for stage in stages_order:
    info = STAGE_PRIORITY[stage]
    words = sorted(all_stage_words.get(stage, set()))
    stages_json[stage] = {
        "priority": info["priority"],
        "label": info["label"],
        "words": words,
        "count": len(words),
    }

word_to_stage_json = {}
for word in sorted(word_first_stage.keys()):
    word_to_stage_json[word] = {
        "first_stage": word_first_stage[word],
        "all_stages": word_all_stages[word],
        "first_priority": word_first_priority[word],
    }

overlap_json = {}
for low_stage in stages_order:
    overlap_json[low_stage] = {}
    for high_stage in stages_order:
        val = overlap_matrix[low_stage].get(high_stage)
        if val is not None:
            overlap_json[low_stage][high_stage] = val

output = {
    "stages": stages_json,
    "word_to_stage": word_to_stage_json,
    "overlap_matrix": overlap_json,
}

outpath = os.path.join(OUT_DIR, "stage_vocab.json")
with open(outpath, "w", encoding="utf-8") as f:
    json.dump(output, f, ensure_ascii=False, indent=2)
print(f"已写入: {outpath}")
print(f"JSON 大小: {os.path.getsize(outpath)/1024:.1f} KB")

# 也输出统计摘要
print(f"\n总计: {len(word_first_stage)} 个去重单词（所有阶段合并）")

# ══════════════════════════════════════════════════════════════
# 问题汇总（用于文档）
# ══════════════════════════════════════════════════════════════
print("\n\n--- 问题汇总 ---")
all_issue_count = 0
for fname, iss in all_issues.items():
    if iss:
        print(f"\n  {fname}: {len(iss)} 个问题")
        all_issue_count += len(iss)
print(f"\n总共 {all_issue_count} 个问题")
